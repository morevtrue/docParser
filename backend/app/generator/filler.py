"""
Динамический генератор выходных DOCX.

Алгоритм:
1. Сканирует шаблон — находит все плейсхолдеры вида [...] в параграфах и таблицах.
2. Ищет каждый плейсхолдер в FIELD_MAP по точному совпадению.
3. Извлекает значение из ExtractedData по source + field.
4. Применяет трансформацию (если указана).
5. Подставляет значение или оставляет пустым + пишет в warnings/unmapped.
"""
import re
from datetime import date
from docx import Document
from ..models.schemas import ExtractedData
from ..mapping.field_map import FIELD_MAP, PRICE_TABLE_COLUMNS

# Индекс маппинга для быстрого поиска по плейсхолдеру
_FIELD_MAP_INDEX: dict[str, dict] = {entry["placeholder"]: entry for entry in FIELD_MAP}

# Регулярка для поиска плейсхолдеров вида [...]
_PLACEHOLDER_RE = re.compile(r"\[[^\[\]]+\]")


# ---------------------------------------------------------------------------
# Трансформации
# ---------------------------------------------------------------------------

def _shorten_name(value: str) -> str:
    """'Иванов Иван Иванович' → 'Иванов И.И.'"""
    parts = value.strip().split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return value


def _format_money(value: str) -> str:
    """'432600.00 RUB' → '432 600,00'"""
    value = value.split()[0]
    try:
        num = float(value.replace(",", "."))
        integer_part = int(num)
        decimal_part = round((num - integer_part) * 100)
        formatted_int = f"{integer_part:,}".replace(",", " ")
        return f"{formatted_int},{decimal_part:02d}"
    except ValueError:
        return value


def _join_bank(sc) -> str:
    """Собирает многострочный блок банковских реквизитов."""
    lines = []
    if sc.checking_account:
        lines.append(sc.checking_account)
    if sc.bank:
        lines.append(sc.bank)
    if sc.correspondent_account:
        lines.append(f"к/с {sc.correspondent_account}")
    if sc.bik:
        lines.append(f"БИК {sc.bik}")
    return "\n".join(lines)


def _join_contact(sc) -> str:
    """Собирает многострочный блок контактного лица."""
    lines = []
    if sc.contact_person:
        lines.append(sc.contact_person)
    if sc.email:
        lines.append(sc.email)
    if sc.phone:
        lines.append(sc.phone)
    return "\n".join(lines)


_MONTHS_RU = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

# Счётчик исходящих номеров за текущий день (сбрасывается при рестарте)
_outgoing_counter: dict[str, int] = {}


def _current_date(_v: str) -> str:
    """Возвращает дату в формате «26» марта 2026 года."""
    today = date.today()
    return f"«{today.day}» {_MONTHS_RU[today.month]} {today.year} года"


def _outgoing_number(_v: str) -> str:
    """Генерирует исходящий номер вида №2603/1 (ДДММ/порядковый)."""
    today = date.today()
    key = today.strftime("%d%m")
    _outgoing_counter[key] = _outgoing_counter.get(key, 0) + 1
    return f"№{key}/{_outgoing_counter[key]}"


TRANSFORMS = {
    "shorten_name":    lambda v, _dto: _shorten_name(v),
    "format_money":    lambda v, _dto: _format_money(v),
    "current_date":    lambda v, _dto: _current_date(v),
    "outgoing_number": lambda v, _dto: _outgoing_number(v),
}


# ---------------------------------------------------------------------------
# Извлечение значения из DTO по маппингу
# ---------------------------------------------------------------------------

def _resolve_value(entry: dict, data: ExtractedData) -> str | None:
    """
    Извлекает значение из ExtractedData по записи FIELD_MAP.
    Возвращает строку или None если данных нет.
    """
    source = entry["source"]
    field = entry["field"]
    transform = entry.get("transform")

    # Генерируемые значения — не из DTO
    if source == "__generated__":
        if transform and transform in TRANSFORMS:
            return TRANSFORMS[transform]("", data) or None
        return None

    source_obj = getattr(data, source, None)
    if source_obj is None:
        return None

    value = getattr(source_obj, field, None)
    if value is None:
        return None

    value = str(value)
    if transform and transform in TRANSFORMS:
        value = TRANSFORMS[transform](value, data)

    return value or None


# ---------------------------------------------------------------------------
# Сканирование плейсхолдеров в шаблоне
# ---------------------------------------------------------------------------

def _scan_placeholders(doc: Document) -> set[str]:
    """Находит все уникальные плейсхолдеры [...] в параграфах и таблицах документа."""
    found = set()

    for para in doc.paragraphs:
        for m in _PLACEHOLDER_RE.finditer(para.text):
            found.add(m.group())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for m in _PLACEHOLDER_RE.finditer(para.text):
                        found.add(m.group())

    return found


# ---------------------------------------------------------------------------
# Замена в параграфах
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> bool:
    """
    Заменяет плейсхолдеры в параграфе, сохраняя форматирование runs.

    Алгоритм:
    1. Сначала пробуем заменить внутри каждого run по отдельности
       (плейсхолдер целиком в одном run — форматирование сохраняется).
    2. Если плейсхолдер разбит между runs — склеиваем все runs параграфа
       в первый run (берём его форматирование), заменяем, очищаем остальные.
    """
    # Шаг 1: замена внутри отдельных runs
    changed = False
    for run in paragraph.runs:
        new_text = run.text
        for placeholder, value in replacements.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
        if new_text != run.text:
            run.text = new_text
            changed = True

    # Шаг 2: проверяем, остались ли незамененные плейсхолдеры (разбиты между runs)
    full_text = paragraph.text
    remaining = {ph: val for ph, val in replacements.items() if ph in full_text}
    if not remaining:
        return changed

    # Склеиваем все runs в первый, применяем замены
    if paragraph.runs:
        merged = full_text
        for placeholder, value in remaining.items():
            merged = merged.replace(placeholder, value)
        paragraph.runs[0].text = merged
        for run in paragraph.runs[1:]:
            run.text = ""
        return True

    return changed


def _replace_in_cell(cell, replacements: dict[str, str]):
    for para in cell.paragraphs:
        _replace_in_paragraph(para, replacements)


# ---------------------------------------------------------------------------
# Основная функция заполнения
# ---------------------------------------------------------------------------

def fill_template(
    template_source,
    data: ExtractedData,
) -> tuple[Document, dict]:
    """
    Заполняет шаблон DOCX данными из ExtractedData через FIELD_MAP.

    Возвращает (Document, log) где log = {
        "found":    [(placeholder, value, source)],
        "warnings": [(placeholder, source)],   # данные не найдены
        "unmapped": [placeholder],              # нет в FIELD_MAP
    }
    """
    doc = Document(template_source)

    # Шаг 1: сканируем плейсхолдеры в шаблоне
    placeholders = _scan_placeholders(doc)

    # Шаг 2: строим словарь замен + лог
    replacements: dict[str, str] = {}
    log: dict = {"found": [], "warnings": [], "unmapped": []}

    for ph in placeholders:
        entry = _FIELD_MAP_INDEX.get(ph)
        if entry is None:
            log["unmapped"].append(ph)
            replacements[ph] = ""   # убираем плейсхолдер из документа
            continue

        value = _resolve_value(entry, data)
        if value:
            replacements[ph] = value
            log["found"].append({"placeholder": ph, "value": value, "source": entry["source"]})
        else:
            replacements[ph] = ""
            log["warnings"].append({"placeholder": ph, "source": entry["source"]})

    # Шаг 3: применяем замены
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    return doc, log


# ---------------------------------------------------------------------------
# Специальная логика для таблицы номенклатуры
# ---------------------------------------------------------------------------

def _set_cell_text(cell, value: str):
    """Очищает все runs в первом параграфе ячейки и записывает значение."""
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = value
    else:
        para.add_run(value)


def _detect_price_output_columns(header_cells: list[str]) -> dict[int, str]:
    """
    Семантически определяет маппинг колонок таблицы номенклатуры шаблона по заголовкам.
    Возвращает {col_idx: field_name}.
    """
    mapping = {}
    for i, h in enumerate(header_cells):
        h_lower = h.lower().strip()
        if any(kw in h_lower for kw in ["наименование", "название", "товар", "продукция"]):
            mapping[i] = "name"
        elif any(kw in h_lower for kw in ["ед. изм", "ед.изм", "единица", "ед."]):
            mapping[i] = "unit"
        elif any(kw in h_lower for kw in ["нмц", "нмц за ед", "нмц, руб"]):
            mapping[i] = "nmc"
        elif any(kw in h_lower for kw in ["цена без ндс", "цена без", "цена за ед"]):
            mapping[i] = "price_without_vat"
        elif any(kw in h_lower for kw in ["количество", "кол-во", "кол."]):
            mapping[i] = "quantity"
        elif any(kw in h_lower for kw in ["сумма без ндс", "сумма без", "итого без", "сумма"]):
            mapping[i] = "total_without_vat"
    return mapping


def fill_price_template(
    template_source,
    data: ExtractedData,
) -> tuple[Document, dict]:
    """
    Заполняет шаблон «Предложение о цене договора».
    Таблица номенклатуры заполняется по PRICE_TABLE_COLUMNS.
    Остальные плейсхолдеры — через стандартный fill_template.
    """
    doc = Document(template_source)

    ct = data.commercial_terms
    cr = data.customer_request

    # НМЦ по индексу позиции из запроса ТКП
    nmc_by_index: dict[int, str] = {}
    if cr and cr.items:
        for item in cr.items:
            try:
                nmc_by_index[int(item.number or "0") - 1] = item.nmc or ""
            except ValueError:
                pass

    # Заполняем таблицу номенклатуры (Таблица 0)
    if ct and ct.items and len(doc.tables) > 0:
        price_table = doc.tables[0]
        if price_table.rows:
            # Семантически определяем колонки по заголовку
            header_cells = [c.text.strip().lower() for c in price_table.rows[0].cells]
            col_map = _detect_price_output_columns(header_cells)

            for i, item in enumerate(ct.items):
                row_idx = i + 1
                if row_idx >= len(price_table.rows):
                    break
                cells = price_table.rows[row_idx].cells

                for col_idx, field in col_map.items():
                    if col_idx >= len(cells):
                        continue
                    if field == "nmc":
                        value = nmc_by_index.get(i, "")
                    elif field == "total_without_vat":
                        raw = getattr(item, field, "") or ""
                        value = _format_money(raw) if raw else ""
                    else:
                        value = getattr(item, field, "") or ""
                    _set_cell_text(cells[col_idx], value)

    # Остальные плейсхолдеры через стандартный механизм
    placeholders = _scan_placeholders(doc)
    replacements: dict[str, str] = {}
    log: dict = {"found": [], "warnings": [], "unmapped": []}

    for ph in placeholders:
        entry = _FIELD_MAP_INDEX.get(ph)
        if entry is None:
            log["unmapped"].append(ph)
            replacements[ph] = ""
            continue
        value = _resolve_value(entry, data)
        if value:
            replacements[ph] = value
            log["found"].append({"placeholder": ph, "value": value, "source": entry["source"]})
        else:
            replacements[ph] = ""
            log["warnings"].append({"placeholder": ph, "source": entry["source"]})

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    return doc, log
